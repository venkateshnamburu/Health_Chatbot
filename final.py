import streamlit as st
import openai
import fitz  # PyMuPDF for PDFs
import pytesseract  # OCR for images
from PIL import Image
import pdfplumber  # Extract tables from PDFs
import docx  # Extract text from Word files
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Pinecone as PineconeVectorStore
from langchain.llms import OpenAI
from langchain.chains import RetrievalQA
from langchain.retrievers import BM25Retriever
from langchain.schema import Document
import tempfile
import os
import threading
from pinecone import Pinecone, ServerlessSpec

# Set API keys
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")  # Ensure it's set in your environment
PINECONE_API_KEY="pcsk_4G6dht_A9TsQnGHt6GZzcQT7MHTv4K9pLyabYuT639Hu6RZoWqdfN5uHJRQzqNZiiYtBBE"
PINECONE_ENV = "us-east-1"

# Initialize OpenAI client
client = openai.OpenAI(api_key=OPENAI_API_KEY)

# Initialize Pinecone client
pc = Pinecone(api_key=PINECONE_API_KEY)

# Check if index exists, create if needed
INDEX_NAME = "chatbot"
if INDEX_NAME not in pc.list_indexes().names():
    pc.create_index(
        name=INDEX_NAME,
        dimension=1536,
        metric="cosine",
        spec=ServerlessSpec(cloud="aws", region=PINECONE_ENV)
    )

# Connect to Pinecone index
embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
vectorstore = PineconeVectorStore.from_existing_index(INDEX_NAME, embeddings)

# ---- Chat History Setup ----
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

# ---- Text Extraction Functions ----
def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text() + "\n"
    return text

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    return pytesseract.image_to_string(image, lang='eng', config='--psm 6')

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_tables_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    clean_row = [str(cell) if cell is not None else "" for cell in row]
                    text += " | ".join(clean_row) + "\n"
    return text

# ---- Chunking & Indexing ----
def chunk_text(text, chunk_size=500, chunk_overlap=100):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_text(text)

def add_documents_to_pinecone(texts):
    vectorstore.add_texts(texts)

def fetch_all_documents():
    query_result = vectorstore.similarity_search("placeholder_query", k=1000)
    return [Document(page_content=item.page_content) for item in query_result]

# ---- Hybrid Retrieval ----
def hybrid_search(query, k=3):
    stored_documents = fetch_all_documents()
    bm25_retriever = BM25Retriever.from_documents(stored_documents)
    bm25_results = bm25_retriever.get_relevant_documents(query, k=k)
    vector_results = vectorstore.similarity_search(query, k=k)
    return bm25_results + vector_results

# ---- RAG-based Response ----
def generate_response(query):
    relevant_docs = hybrid_search(query)
    context = "\n".join([doc.page_content for doc in relevant_docs])
    
    if not context.strip():
        return "Sorry, I couldn't find relevant medical information."

    prompt = f"""
    You are a healthcare assistant. Answer the question based on the context:
    Context: {context}
    Question: {query}
    """

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[{"role": "system", "content": prompt}]
    )

    return response.choices[0].message.content

# ---- File Processing ----
def process_pdf(file_path, filename):
    st.write(f"Processing PDF: {filename}")
    pdf_text = extract_text_from_pdf(file_path) + extract_tables_from_pdf(file_path)
    chunks = chunk_text(pdf_text)
    add_documents_to_pinecone(chunks)
    st.success(f"Processed {filename}")

# ---- Streamlit UI ----
def main():
    st.title("Healthcare AI Chatbot")

    # Upload files
    uploaded_files = st.file_uploader("Upload PDFs, Images, or Word files", accept_multiple_files=True, type=["pdf", "png", "jpg", "jpeg", "docx"])
    if uploaded_files:
        st.write("Processing uploaded files...")
        for uploaded_file in uploaded_files:
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(uploaded_file.read())
                file_path = temp_file.name
                
                if uploaded_file.type == "application/pdf":
                    threading.Thread(target=process_pdf, args=(file_path, uploaded_file.name)).start()
                elif uploaded_file.type in ["image/png", "image/jpeg"]:
                    st.write(f"Processing Image: {uploaded_file.name}")
                    image_text = extract_text_from_image(file_path)
                    chunks = chunk_text(image_text)
                    add_documents_to_pinecone(chunks)
                elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    st.write(f"Processing Word File: {uploaded_file.name}")
                    doc_text = extract_text_from_docx(file_path)
                    chunks = chunk_text(doc_text)
                    add_documents_to_pinecone(chunks)
        st.success("All files processed and stored in the database!")

    # User Query Section
    user_query = st.text_input("Ask a medical question:")

    if st.button("Get Answer"):
        if user_query.strip():
            with st.spinner("Generating response..."):
                response = generate_response(user_query)

            # Store conversation in session state
            st.session_state.chat_history.append(("User", user_query))
            st.session_state.chat_history.append(("AI", response))

            # Display updated chat history
            st.subheader("Chat History:")
            for role, text in st.session_state.chat_history:
                st.write(f"**{role}:** {text}")

        else:
            st.warning("Please enter a valid question.")

    # Query Completion Button (Clears Chat History)
    if st.button("Query Completed"):
        st.session_state.chat_history = []
        st.success("Chat history cleared.")

if __name__ == "__main__":
    main()
